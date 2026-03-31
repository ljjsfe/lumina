"""DOCX file profiling via python-docx."""

from __future__ import annotations

import os

from ..core.types import ManifestEntry


def read_docx(file_path: str) -> ManifestEntry:
    """Profile a DOCX file into a ManifestEntry."""
    size = os.path.getsize(file_path)

    try:
        from docx import Document
    except ImportError:
        return ManifestEntry(
            file_path=file_path, file_type="docx", size_bytes=size,
            summary={"error": "python-docx not installed"},
        )

    try:
        doc = Document(file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(text_parts)

        # Key terms from bold runs
        key_terms = set()
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold and run.text.strip():
                    key_terms.add(run.text.strip())

        return ManifestEntry(
            file_path=file_path, file_type="docx", size_bytes=size,
            summary={
                "headings": headings[:20],
                "key_terms": sorted(key_terms)[:20],
                "text_preview": full_text[:1000],
                "paragraph_count": len(text_parts),
            },
        )
    except Exception as e:
        return ManifestEntry(
            file_path=file_path, file_type="docx", size_bytes=size,
            summary={"error": str(e)},
        )
